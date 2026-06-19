# -*- coding: utf-8 -*-
"""Word (.docx) export for documented modules, user-manual style.

Generates a complete Russian-language user manual in Word format:

* Title page (system name, "Руководство пользователя", version, developer, city/year)
* Real Word TOC field ("ОГЛАВЛЕНИЕ")
* 1. Введение (1.1 Категории, 1.2 Область, 1.3 Назначение, 1.4 Соглашения)
* 2. Содержание документа (2.1 Назначение, 2.2 Материалы, 2.3 Подготовка)
* 3. Жизненный цикл и состояния — optional, driven by doc.module.workflow_states
* 4. Наследуемые модели и поля — optional, driven by doc.module.inherited_model_ids
* 5. Интеграции — optional, driven by doc.module.integration_ids
* 6. Список функций — one block per function
* 7. Аналитика и экспорт — optional, driven by doc.module.analytics_ids
* 8. Литература
* 9. Словарь терминов

Sections 3-5 and 7 are SKIPPED when the corresponding data is absent,
so existing modules that have no workflow/integration data are not affected.

All new sections are driven by generic doc.module relationship fields —
NOT hardcoded for dpf_events — so any addon benefits automatically.

``python-docx`` is loaded lazily so the addon still installs without it;
the Word button then raises a clear, actionable error.
"""
import base64
import io
import logging

from odoo import _, api, models
from odoo.exceptions import UserError

_logger = logging.getLogger(__name__)

try:
    import docx  # python-docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:  # pragma: no cover
    HAS_DOCX = False

BODY_FONT = "Times New Roman"
BODY_SIZE = 12


def _grey():
    return RGBColor(0x80, 0x80, 0x80)


def _red():
    return RGBColor(0xC0, 0x00, 0x00)


def _blue():
    return RGBColor(0x17, 0x56, 0xAB)


class DocWordExport(models.AbstractModel):
    _name = "doc.word.export"
    _description = "DPF Docs - Word Export Service"

    # ------------------------------------------------------------------
    # Public entry point
    # ------------------------------------------------------------------
    @api.model
    def build_docx(self, doc_modules):
        """Return raw .docx bytes for all given doc.module records."""
        if not HAS_DOCX:
            raise UserError(_(
                "Пакет 'python-docx' не установлен на сервере.\n"
                "Попросите администратора выполнить:  pip install python-docx"
            ))

        self._renumber_functions(doc_modules)

        document = docx.Document()
        self._apply_base_styles(document)

        primary = doc_modules[:1]
        self._add_cover(document, primary)
        self._add_toc(document)
        self._add_intro_section(document, primary)
        self._add_content_section(document, primary)

        # --- Optional sections (skipped when data absent) ---
        self._add_workflow_section(document, primary)
        self._add_inherited_models_section(document, primary)
        self._add_integrations_section(document, primary)

        self._add_functions_section(document, doc_modules)

        self._add_analytics_section(document, primary)
        self._add_bibliography_section(document, primary)
        self._add_glossary_section(document, primary)

        buf = io.BytesIO()
        document.save(buf)
        return buf.getvalue()

    # ------------------------------------------------------------------
    # Function renumbering
    # ------------------------------------------------------------------
    @api.model
    def _renumber_functions(self, doc_modules):
        counter = 0
        for doc_module in doc_modules:
            funcs = doc_module.function_ids.sorted(
                key=lambda f: ((f.sequence or 999999), f.id)
            )
            for func in funcs:
                counter += 1
                if func.number != counter:
                    func.write({'number': counter})

    # ------------------------------------------------------------------
    # Styling
    # ------------------------------------------------------------------
    def _apply_base_styles(self, document):
        style = document.styles["Normal"]
        style.font.name = BODY_FONT
        style.font.size = Pt(BODY_SIZE)
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:cs"), BODY_FONT)
        rfonts.set(qn("w:eastAsia"), BODY_FONT)

    # ------------------------------------------------------------------
    # Cover page
    # ------------------------------------------------------------------
    def _add_cover(self, document, modules):
        m = modules[:1]
        system_name = (m.system_name if m else None) or (
            'Система "%s"' % (m.name if m else "Odoo")
        )
        platform = (m.platform_version if m else None) or "Odoo 19"
        version = (m.manual_version if m else None) or "1.0"
        developer = (m.developer if m else None) or ""
        city_year = (m.city_year if m else None) or ""

        for _ in range(7):
            document.add_paragraph()

        self._centered(document, system_name, size=18)
        self._centered(document, 'на базе платформы "%s"' % platform, size=16)
        document.add_paragraph()
        self._centered(document, "Руководство пользователя", bold=True, size=18)
        self._centered(document, "Версия %s" % version, size=14)

        for _ in range(4):
            document.add_paragraph()

        if developer:
            self._centered(document, "Разработчик: %s" % developer, size=12)

        for _ in range(6):
            document.add_paragraph()

        if city_year:
            self._centered(document, city_year, size=12)

        document.add_page_break()

    def _centered(self, document, text, bold=False, size=12):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.name = BODY_FONT
        return p

    # ------------------------------------------------------------------
    # Table of contents
    # ------------------------------------------------------------------
    def _add_toc(self, document):
        h = document.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h.add_run("ОГЛАВЛЕНИЕ")
        run.bold = True
        run.font.size = Pt(14)
        document.add_paragraph()

        p = document.add_paragraph()
        r = p.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = r'TOC \o "1-3" \h \z \u'
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        hint = OxmlElement("w:t")
        hint.text = (
            "Оглавление обновляется автоматически: "
            "правый клик → «Обновить поле»."
        )
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        r._r.extend([fld_begin, instr, fld_sep, hint, fld_end])
        document.add_page_break()

    # ------------------------------------------------------------------
    # Heading helpers
    # ------------------------------------------------------------------
    def _heading(self, document, text, level):
        p = document.add_heading("", level=level)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.name = BODY_FONT
        r.font.size = Pt(14 if level == 1 else 12)
        return p

    def _paragraph(self, document, text):
        if text and text.strip():
            document.add_paragraph(text.strip())

    def _bullets(self, document, text):
        """Each non-empty line becomes a bullet."""
        for line in (text or "").splitlines():
            line = line.strip()
            if line:
                document.add_paragraph(line, style="List Bullet")

    def _numbered_lines(self, document, text):
        """Each non-empty line becomes a numbered paragraph starting at 1."""
        for i, line in enumerate(
            [l.strip() for l in (text or "").splitlines() if l.strip()], 1
        ):
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            p.add_run("%d. %s" % (i, line))

    def _simple_table(self, document, headers, rows, col_widths=None):
        """
        Insert a simple bordered table.

        Parameters
        ----------
        headers    : list[str]
        rows       : list[list[str]]
        col_widths : list[float] | None — Inches per column; auto if None
        """
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        hdr_row = table.rows[0]
        for i, hdr in enumerate(headers):
            cell = hdr_row.cells[i]
            cell.text = hdr
            cell.paragraphs[0].runs[0].bold = True
            if col_widths:
                cell.width = Inches(col_widths[i])
        for row_data in rows:
            row = table.add_row()
            for i, val in enumerate(row_data):
                row.cells[i].text = str(val or "")
                if col_widths:
                    row.cells[i].width = Inches(col_widths[i])
        document.add_paragraph()
        return table

    # ------------------------------------------------------------------
    # 1. Введение
    # ------------------------------------------------------------------
    def _add_intro_section(self, document, modules):
        m = modules[:1]
        self._heading(document, "1. Введение", 1)

        self._heading(document, "1.1. Описание категорий пользователей", 2)
        self._bullets(document, m.intro_user_categories if m else "")

        self._heading(document, "1.2. Область применения", 2)
        self._bullets(document, m.intro_scope if m else "")

        self._heading(document, "1.3. Назначение документа", 2)
        self._paragraph(document, m.intro_purpose if m else "")

        self._heading(document, "1.4. Соглашения", 2)
        self._bullets(document, m.intro_conventions if m else "")

    # ------------------------------------------------------------------
    # 2. Содержание документа
    # ------------------------------------------------------------------
    def _add_content_section(self, document, modules):
        m = modules[:1]
        self._heading(document, "2. Содержание документа", 1)

        self._heading(document, "2.1. Назначение", 2)
        self._paragraph(document, m.content_purpose if m else "")

        self._heading(document, "2.2. Необходимые материалы", 2)
        self._bullets(document, m.content_materials if m else "")

        self._heading(document, "2.3. Подготовка к работе", 2)
        self._numbered_lines(document, m.content_preparation if m else "")

    # ------------------------------------------------------------------
    # 3. Жизненный цикл и состояния (OPTIONAL)
    # ------------------------------------------------------------------
    def _add_workflow_section(self, document, modules):
        """
        Render a workflow / lifecycle section driven by doc.module.workflow_state_ids.

        The section is SKIPPED entirely when no workflow states are defined,
        so modules without a state machine are not affected.

        Expected model: doc.workflow.state with fields:
            name           (char)  — state technical name / key
            label          (char)  — human-readable label
            description    (text)  — what this state means for the user
            transitions    (text)  — comma/newline separated list of next states
            button_label   (char)  — UI button that triggers the transition (optional)
        """
        m = modules[:1]
        if not m:
            return
        states = getattr(m, 'workflow_state_ids', None)
        if not states:
            return

        self._heading(document, "3. Жизненный цикл объекта", 1)
        self._paragraph(
            document,
            "В данном разделе описаны все возможные состояния объекта и переходы "
            "между ними. Состояния управляются специальными кнопками в форме записи."
        )

        headers = ["Состояние", "Метка", "Описание", "Переходы", "Кнопка"]
        col_widths = [1.0, 1.1, 2.4, 1.5, 1.0]
        rows = []
        for state in states:
            rows.append([
                state.name or "",
                state.label or "",
                state.description or "",
                state.transitions or "",
                state.button_label or "",
            ])
        self._simple_table(document, headers, rows, col_widths)

    # ------------------------------------------------------------------
    # 4. Наследуемые модели и поля (OPTIONAL)
    # ------------------------------------------------------------------
    def _add_inherited_models_section(self, document, modules):
        """
        Render an inherited-models section driven by doc.module.inherited_model_ids.

        Covers scenarios where the addon extends a base model via _inherit
        without creating its own top-level menu entry, which means the standard
        introspector misses those fields entirely.

        The section is SKIPPED when no inherited model records exist.

        Expected model: doc.inherited.model with fields:
            base_model     (char)  — original model name, e.g. 'event.event'
            description    (text)  — what this extension adds
            field_ids      (o2m)   — doc.inherited.field records:
                field_name (char), field_type (char), description (text),
                is_required (bool), is_computed (bool)
        """
        m = modules[:1]
        if not m:
            return
        inh_models = getattr(m, 'inherited_model_ids', None)
        if not inh_models:
            return

        self._heading(document, "4. Расширения базовых моделей", 1)
        self._paragraph(
            document,
            "Данный модуль расширяет следующие базовые модели Odoo, добавляя "
            "к ним новые поля и бизнес-логику."
        )

        for idx, inh in enumerate(inh_models, 1):
            self._heading(
                document,
                "4.%d. %s" % (idx, inh.base_model or "Неизвестная модель"),
                2
            )
            if inh.description:
                self._paragraph(document, inh.description)

            fields = getattr(inh, 'field_ids', None)
            if fields:
                headers = ["Поле", "Тип", "Обязательное", "Вычисляемое", "Описание"]
                col_widths = [1.5, 1.0, 0.9, 0.9, 2.7]
                rows = []
                for fld in fields:
                    rows.append([
                        fld.field_name or "",
                        fld.field_type or "",
                        "Да" if getattr(fld, 'is_required', False) else "Нет",
                        "Да" if getattr(fld, 'is_computed', False) else "Нет",
                        fld.description or "",
                    ])
                self._simple_table(document, headers, rows, col_widths)

    # ------------------------------------------------------------------
    # 5. Интеграции (OPTIONAL)
    # ------------------------------------------------------------------
    def _add_integrations_section(self, document, modules):
        """
        Render an integrations section driven by doc.module.integration_ids.

        Covers external services (MinIO, RabbitMQ, Auth service, etc.) that
        live in the services/ layer and are invisible to the ORM introspector.

        The section is SKIPPED when no integration records exist.

        Expected model: doc.integration with fields:
            name           (char)  — service name, e.g. 'MinIO', 'RabbitMQ'
            protocol       (char)  — HTTP / AMQP / SMTP / etc.
            purpose        (text)  — what this integration does for the user
            config_hint    (text)  — how to configure / enable it (optional)
        """
        m = modules[:1]
        if not m:
            return
        integrations = getattr(m, 'integration_ids', None)
        if not integrations:
            return

        self._heading(document, "5. Внешние интеграции", 1)
        self._paragraph(
            document,
            "Модуль взаимодействует со следующими внешними сервисами. "
            "Для корректной работы соответствующих функций необходимо их настроить."
        )

        headers = ["Сервис", "Протокол", "Назначение", "Настройка"]
        col_widths = [1.2, 0.9, 2.7, 2.2]
        rows = []
        for itg in integrations:
            rows.append([
                itg.name or "",
                itg.protocol or "",
                itg.purpose or "",
                itg.config_hint or "",
            ])
        self._simple_table(document, headers, rows, col_widths)

    # ------------------------------------------------------------------
    # 6. Список функций
    # ------------------------------------------------------------------
    def _add_functions_section(self, document, doc_modules):
        self._heading(document, "6. Список функций", 1)

        figure_counter = [0]
        func_counter = [0]

        sub = 0
        for doc_module in doc_modules:
            sub += 1
            mod_name = doc_module.name or doc_module.technical_name
            self._heading(document, "6.%d. %s" % (sub, mod_name), 2)

            funcs = list(doc_module.function_ids.sorted(
                key=lambda f: ((f.sequence or 999999), f.id)
            ))
            if not funcs:
                p = document.add_paragraph()
                r = p.add_run("Функции для данного модуля не сформированы.")
                r.font.color.rgb = _grey()
                r.italic = True
                continue

            extra_descs = {}
            orphan_projects = []
            last_auto = None

            for func in funcs:
                src = (getattr(func, 'source', 'auto') or 'auto')
                if src != 'project':
                    last_auto = func
                else:
                    desc = (func.description or '').strip()
                    if not desc:
                        continue
                    if last_auto is not None:
                        extra_descs.setdefault(last_auto.id, []).append(desc)
                    else:
                        orphan_projects.append(func)

            for func in funcs:
                src = (getattr(func, 'source', 'auto') or 'auto')
                if src == 'project':
                    continue
                func_counter[0] += 1
                appended = extra_descs.get(func.id, [])
                self._add_function(
                    document, func,
                    func_num=func_counter[0],
                    figure_counter=figure_counter,
                    extra_descriptions=appended,
                )

            for func in orphan_projects:
                self._add_orphan_project_function(document, func)

    def _add_function(self, document, func, func_num, figure_counter,
                      extra_descriptions=None):
        extra_descriptions = extra_descriptions or []

        title_p = document.add_paragraph()
        title_p.paragraph_format.space_before = Pt(10)
        r = title_p.add_run("Функция %d: %s." % (func_num, func.name or ""))
        r.bold = True
        r.font.size = Pt(12)

        desc_parts = []
        if func.description and func.description.strip():
            desc_parts.append(func.description.strip())
        for extra in extra_descriptions:
            if extra:
                desc_parts.append(extra)

        if desc_parts:
            p = document.add_paragraph()
            lbl = p.add_run("Описание: ")
            lbl.bold = True
            p.add_run("\n".join(desc_parts))

        if func.requirements and func.requirements.strip():
            p = document.add_paragraph()
            lbl = p.add_run("Требования: ")
            lbl.bold = True
            val = p.add_run(func.requirements.strip())
            val.font.color.rgb = _red()

        steps = func.step_lines() if hasattr(func, "step_lines") else []
        if steps:
            p = document.add_paragraph()
            p.add_run("Порядок выполнения:").bold = True
            for idx, step in enumerate(steps, 1):
                sp = document.add_paragraph()
                sp.paragraph_format.left_indent = Pt(18)
                sp.add_run("%d. %s" % (idx, step))

        if func.screenshot:
            figure_counter[0] = self._embed_screenshot(
                document, func, figure_counter[0]
            )
        else:
            figure_counter[0] = self._add_screenshot_placeholder(
                document, func, figure_counter[0]
            )

        self._labelled(document, "Результат:", func.result)

        sep = document.add_paragraph()
        sep.paragraph_format.space_after = Pt(6)

    def _add_orphan_project_function(self, document, func):
        desc = (func.description or '').strip()
        if not desc:
            return
        p = document.add_paragraph()
        r = p.add_run("%s: %s" % (func.name or "Дополнение", desc))
        r.italic = True
        r.font.color.rgb = _grey()
        r.font.size = Pt(11)
        sep = document.add_paragraph()
        sep.paragraph_format.space_after = Pt(4)

    def _labelled(self, document, label, value):
        if not value or not value.strip():
            return
        p = document.add_paragraph()
        r = p.add_run(label + " ")
        r.bold = True
        p.add_run(value.strip())

    def _embed_screenshot(self, document, func, figure):
        try:
            stream = io.BytesIO(base64.b64decode(func.screenshot))
            document.add_picture(stream, width=Inches(5.5))
            document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            figure += 1
            cap_p = document.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_text = func.screenshot_caption or func.name or ""
            cap_r = cap_p.add_run("Рис.%d %s" % (figure, caption_text))
            cap_r.font.size = Pt(10)
            cap_r.italic = True
        except Exception as exc:  # pragma: no cover
            _logger.warning(
                "Could not embed screenshot for function %s: %s", func.id, exc
            )
        return figure

    def _add_screenshot_placeholder(self, document, func, figure):
        figure += 1
        name = func.name or "экрана"
        caption_text = func.screenshot_caption or name

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(
            "📌 [Здесь должен быть скриншот экрана «%s»]" % name
        )
        r.font.color.rgb = _grey()
        r.italic = True
        r.font.size = Pt(11)

        cap_p = document.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_r = cap_p.add_run("Рис.%d %s" % (figure, caption_text))
        cap_r.font.size = Pt(10)
        cap_r.italic = True
        cap_r.font.color.rgb = _grey()

        return figure

    # ------------------------------------------------------------------
    # 7. Аналитика и экспорт (OPTIONAL)
    # ------------------------------------------------------------------
    def _add_analytics_section(self, document, modules):
        """
        Render an analytics section driven by doc.module.analytic_field_ids
        and doc.module.export_action_ids.

        Covers computed KPI fields and export buttons that ORM introspector
        misses because they are not user-input fields (compute=True, store=False).

        Section is SKIPPED when neither analytic_field_ids nor export_action_ids
        are present, keeping the document clean for simple modules.

        Expected models:
          doc.analytic.field:
            name (char), description (text), formula_hint (text)
          doc.export.action:
            name (char), format (char), description (text)
        """
        m = modules[:1]
        if not m:
            return
        analytic_fields = getattr(m, 'analytic_field_ids', None)
        export_actions = getattr(m, 'export_action_ids', None)

        if not analytic_fields and not export_actions:
            return

        self._heading(document, "7. Аналитика и экспорт", 1)

        if analytic_fields:
            self._heading(document, "7.1. Вычисляемые показатели", 2)
            self._paragraph(
                document,
                "Следующие показатели рассчитываются автоматически на основании "
                "данных в системе и недоступны для ручного редактирования."
            )
            headers = ["Показатель", "Описание", "Формула / источник"]
            col_widths = [1.5, 2.5, 3.0]
            rows = []
            for af in analytic_fields:
                rows.append([
                    af.name or "",
                    af.description or "",
                    af.formula_hint or "",
                ])
            self._simple_table(document, headers, rows, col_widths)

        if export_actions:
            self._heading(document, "7.2. Экспорт данных", 2)
            self._paragraph(
                document,
                "Для выгрузки отчётов используйте следующие действия, "
                "доступные в форме записи."
            )
            headers = ["Действие", "Формат", "Описание"]
            col_widths = [1.8, 0.8, 4.4]
            rows = []
            for ea in export_actions:
                rows.append([
                    ea.name or "",
                    ea.format or "",
                    ea.description or "",
                ])
            self._simple_table(document, headers, rows, col_widths)

    # ------------------------------------------------------------------
    # 8. Литература / 9. Словарь
    # ------------------------------------------------------------------
    def _add_bibliography_section(self, document, modules):
        m = modules[:1]
        document.add_page_break()
        self._heading(document, "8. Литература", 1)
        self._bullets(document, m.bibliography if m else "")

    def _add_glossary_section(self, document, modules):
        m = modules[:1]
        self._heading(document, "9. Словарь терминов", 1)
        self._bullets(document, m.glossary if m else "")
